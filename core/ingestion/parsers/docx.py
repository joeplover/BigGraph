from pathlib import Path
from docx import Document as DocxDocument
from core.ingestion.parser_base import ParsedDocument, ParsedSection


class DocxParser:
    supported_extensions = {".docx"}
    def parse(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(str(path))
        sections, md, cur, buf = [], [], path.stem, []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower()
            if style.startswith("heading") or "标题" in style:
                if buf:
                    sections.append(ParsedSection(heading_path=cur, text="\n".join(buf)))
                    buf = []
                cur = f"{path.stem} / {text}"
                md.append(f"## {text}")
            else:
                buf.append(text)
                md.append(text)
        if buf:
            sections.append(ParsedSection(heading_path=cur, text="\n".join(buf)))
        for table in doc.tables:
            rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
            if rows:
                cols = max(len(r) for r in rows)
                tbl = "\n".join(
                    ["| " + " | ".join(rows[0] + [""] * (cols - len(rows[0]))) + " |",
                     "| " + " | ".join(["---"] * cols) + " |"] +
                    ["| " + " | ".join(r + [""] * (cols - len(r))) + " |" for r in rows[1:]]
                )
                md.append(tbl)
                sections.append(ParsedSection(heading_path=f"{cur} / 表格", text=tbl, content_type="table"))
        md_text = "\n\n".join(md)
        return ParsedDocument(title=path.stem, source_path=str(path), raw_text=md_text, normalized_markdown=md_text, sections=sections or [ParsedSection(heading_path=path.stem, text=md_text)], metadata={"parser": "docx"})
